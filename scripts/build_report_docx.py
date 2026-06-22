#!/usr/bin/env python3
"""Build a Word report from Thoracic Weekly Academy structured content JSON."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


DISEASE_ORDER = ["肺癌", "纵隔肿瘤", "食管癌", "气胸、胸部外伤、肋骨骨折、胸壁畸形"]
DISEASE_COUNT_LABELS = {
    "lung_cancer": "肺癌",
    "mediastinal_tumor": "纵隔肿瘤",
    "esophageal_cancer": "食管癌",
    "pneumothorax": "气胸",
    "pneumothorax_chest_trauma_rib_fracture_chest_wall_deformity": "气胸、胸部外伤、肋骨骨折、胸壁畸形",
    "mediastinal_tumor_supplement": "纵隔肿瘤补充检索",
}
TYPE_ORDER = ["临床研究", "人工智能/机器学习相关研究", "其他基础研究"]


def normalize_journal(value):
    value = (value or "").lower()
    value = re.sub(r"[^a-z0-9]+", " ", value)
    return " ".join(value.split())


def load_journal_metrics(skill_dir):
    metrics_path = skill_dir / "references" / "journal_metrics.json"
    metrics = json.loads(metrics_path.read_text())
    by_key = {}
    for journal in metrics["journals"]:
        keys = [journal.get("journal", "")]
        keys.extend(journal.get("pubmed_journal_terms") or [])
        for key in keys:
            normalized = normalize_journal(key)
            if normalized:
                by_key[normalized] = journal
    return by_key


def journal_metric(item, metrics_by_key):
    keys = [
        item.get("journal"),
        item.get("journal_abbr"),
    ]
    for key in keys:
        normalized = normalize_journal(key)
        if normalized in metrics_by_key:
            return metrics_by_key[normalized]
    return None


def journal_metric_label(item, metrics_by_key):
    metric = journal_metric(item, metrics_by_key)
    if not metric:
        return "未缓存"
    parts = [
        f"JCR {metric.get('jcr_quartile', '')}".strip(),
        f"IF {metric.get('impact_factor', '')}".strip(),
    ]
    if metric.get("new_talent_quartile"):
        parts.append(f"新锐 {metric['new_talent_quartile']}")
    return "\n".join(part for part in parts if part and not part.endswith(" "))


def set_font(run, size=10.5, bold=None, italic=None, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    for style_name in ["Normal", "Body Text"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
    for style_name, size in [("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=60, start=80, bottom=60, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def format_table(table, widths_cm=None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths_cm:
        for row in table.rows:
            for idx, width in enumerate(widths_cm):
                row.cells[idx].width = Cm(width)
    for row_idx, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:cantSplit"))
        if row_idx == 0:
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            tr_pr.append(header)
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_idx == 0:
                set_cell_shading(cell, "D9EAF7")
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    set_font(run, size=8.5, bold=(row_idx == 0))


def para(doc, text="", size=10.5, bold=False, italic=False, first_indent=False):
    text = compact_citations(text)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.7)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=18, bold=True, color="17365D")


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        set_font(run, size={1: 14.5, 2: 12.5, 3: 11}[level], bold=True, color="17365D")


def compact_citations(text):
    def repl(match):
        nums = [int(x) for x in re.findall(r"\d+", match.group(0))]
        if not nums:
            return match.group(0)
        ranges = []
        start = prev = nums[0]
        for num in nums[1:]:
            if num == prev + 1:
                prev = num
                continue
            ranges.append(f"{start}-{prev}" if start != prev else str(start))
            start = prev = num
        ranges.append(f"{start}-{prev}" if start != prev else str(start))
        return "[" + ",".join(ranges) + "]"

    return re.sub(r"(?:\[\d+\])+", repl, text or "")


def authors_ama(authors):
    authors = authors or []
    if len(authors) > 6:
        return ", ".join(authors[:3]) + ", et al"
    return ", ".join(authors)


def ama_ref(item):
    authors = authors_ama(item.get("authors"))
    title = item.get("title", "").rstrip(".") + "."
    journal = item.get("journal_abbr") or item.get("journal") or ""
    pubdate = item.get("pubdate", "")
    doi = item.get("doi", "")
    parts = [authors + "." if authors else "", title, f"{journal}.", f"Published online {pubdate}." if pubdate else ""]
    if doi:
        parts.append(f"doi:{doi}")
    if item.get("pmid"):
        parts.append(f"PMID: {item['pmid']}.")
    return " ".join(p for p in parts if p).replace("..", ".")


def section_map(content):
    mapped = {}
    for section in content.get("review_sections", []):
        mapped[(section["disease"], section["type"])] = section
    return mapped


def add_item_table(doc, items, metrics_by_key):
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for idx, header in enumerate(["引用", "题名 / 期刊", "JCR/IF/新锐", "研究要点", "临床或方法学提示", "PMID"]):
        table.rows[0].cells[idx].text = header
    for item in items:
        cells = table.add_row().cells
        cells[0].text = f"[{item['ref_no']}]"
        cells[1].text = f"{item.get('title','')}\n{item.get('journal_abbr') or item.get('journal','')}; {item.get('pubdate','')}"
        cells[2].text = journal_metric_label(item, metrics_by_key)
        cells[3].text = compact_citations(item.get("summary", ""))
        cells[4].text = compact_citations(item.get("note", ""))
        p = cells[5].paragraphs[0]
        p.text = ""
        pmid = item.get("pmid", "")
        if pmid:
            add_hyperlink(p, pmid, f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/")
    format_table(table, [1.0, 4.5, 1.6, 4.6, 4.0, 1.5])


def add_exclusion_table(doc, records):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for idx, header in enumerate(["PMID", "命中来源", "题名 / 期刊", "排除原因"]):
        table.rows[0].cells[idx].text = header
    for rec in records:
        cells = table.add_row().cells
        p = cells[0].paragraphs[0]
        p.text = ""
        pmid = rec.get("pmid", "")
        if pmid:
            add_hyperlink(p, pmid, f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/")
        cells[1].text = rec.get("hit_source", "")
        cells[2].text = f"{rec.get('title','')}\n{rec.get('journal','')}; {rec.get('pubdate','')}"
        cells[3].text = rec.get("reason", "")
    format_table(table, [1.9, 3.0, 7.7, 5.1])


def add_jcr_table(doc, content, skill_dir):
    metrics_path = skill_dir / "references" / "journal_metrics.json"
    metrics = json.loads(metrics_path.read_text())
    by_name = {j["journal"]: j for j in metrics["journals"]}
    used = sorted({item.get("journal") for s in content.get("review_sections", []) for item in s.get("items", []) if item.get("journal")})
    doc.add_page_break()
    add_heading(doc, "文末：本报告涉及期刊的JCR分区、影响因子与新锐分区", 1)
    para(doc, "指标来源：Skill内置缓存；JCR与影响因子来自用户提供的JCR 2025 Excel工作簿，新锐分区来自用户提供的2026新锐分区.xlsx。", size=9.5, italic=True)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for idx, header in enumerate(["期刊", "JCR分区", "最新影响因子", "新锐分区", "指标年份"]):
        table.rows[0].cells[idx].text = header
    for journal in used:
        metric = by_name.get(journal)
        if not metric:
            continue
        cells = table.add_row().cells
        cells[0].text = metric["journal"]
        cells[1].text = metric.get("jcr_quartile", "")
        cells[2].text = str(metric.get("impact_factor", ""))
        cells[3].text = metric.get("new_talent_quartile") or ""
        cells[4].text = metric.get("metric_year", "")
    format_table(table, [5.9, 1.8, 2.3, 2.0, 4.6])


def build(content, out_path, include_jcr=False, skill_dir=None):
    doc = Document()
    style_doc(doc)
    skill_dir = skill_dir or Path(__file__).resolve().parents[1]
    metrics_by_key = load_journal_metrics(skill_dir)
    add_title(doc, content.get("title") or "PubMed上一个自然周（epdat）胸外科相关文献检索报告")
    para(doc, f"生成日期：{content.get('generated_date','')}；检索限定：Electronic Date of Publication [epdat] = {content.get('epdat_start','')} 至 {content.get('epdat_end','')}（周一至周日）。")
    para(doc, "疾病范围：肺癌、纵隔肿瘤、食管癌，以及气胸、胸部外伤、肋骨骨折、胸壁畸形。期刊范围：按预设高影响综合医学、肿瘤、胸外科/呼吸、消化/食管、基础科学和医学AI期刊列表逐项限定。")
    para(doc, "纳入原则：仅纳入原始研究、系统综述/荟萃分析或与目标疾病高度相关的转化/机制综述；新闻、研究摘要、作者反思、回复信、评论，以及仅被关键词误命中但非目标疾病的记录不进入正文。")

    add_heading(doc, "一、检索与审计结果", 1)
    counts = content.get("disease_counts", {})
    count_text = ", ".join(f"{DISEASE_COUNT_LABELS.get(k, k)}: {v}" for k, v in counts.items())
    para(doc, f"严格疾病检索命中数：{count_text}。")
    para(doc, f"合并去重后共{content.get('unique_records', 0)}条记录；正文纳入{content.get('included_count', 0)}条，排除{content.get('excluded_count', 0)}条。所有排除记录及原因见附录。")

    sections = section_map(content)
    add_heading(doc, "二、按疾病与研究类型整理", 1)
    for disease in DISEASE_ORDER:
        add_heading(doc, disease, 2)
        for typ in TYPE_ORDER:
            add_heading(doc, typ, 3)
            sec = sections.get((disease, typ), {"review_paragraphs": [], "items": [], "empty_note": "本类未检索到符合条件的文献。"})
            for text in sec.get("review_paragraphs", []):
                para(doc, text, first_indent=True)
            if sec.get("items"):
                add_item_table(doc, sec["items"], metrics_by_key)
            else:
                para(doc, sec.get("empty_note") or "本类未检索到符合条件的文献。", size=10)

    if content.get("overall_interpretation"):
        add_heading(doc, "三、总体解读", 1)
        for text in content["overall_interpretation"]:
            para(doc, text)

    included = sorted(
        [item for sec in content.get("review_sections", []) for item in sec.get("items", [])],
        key=lambda x: int(x["ref_no"]),
    )
    add_heading(doc, "四、参考文献", 1)
    for item in included:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.45)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{item['ref_no']}. {ama_ref(item)} ")
        set_font(run, size=9.5)
        if item.get("pmid"):
            add_hyperlink(p, "PubMed", f"https://pubmed.ncbi.nlm.nih.gov/{item['pmid']}/")

    add_heading(doc, "附录：排除/未纳入正文记录核对表", 1)
    add_exclusion_table(doc, content.get("excluded_records", []))
    if include_jcr and skill_dir:
        add_jcr_table(doc, content, skill_dir)
    doc.save(out_path)


def validate(content):
    included = [item for sec in content.get("review_sections", []) for item in sec.get("items", [])]
    nums = sorted(int(item["ref_no"]) for item in included)
    expected = list(range(1, len(nums) + 1))
    errors = []
    if nums != expected:
        errors.append(f"Reference numbers are not continuous: {nums}")
    pmids = [item.get("pmid") for item in included if item.get("pmid")]
    dupes = sorted({pmid for pmid in pmids if pmids.count(pmid) > 1})
    if dupes:
        errors.append(f"Duplicate included PMIDs: {dupes}")
    body = json.dumps(content.get("review_sections", []), ensure_ascii=False)
    cited = sorted({int(x) for x in re.findall(r"\[(\d+)\]", body)})
    missing = sorted(set(expected) - set(cited))
    if missing:
        errors.append(f"References not cited in review/table content: {missing}")
    return errors


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("content_json")
    parser.add_argument("out_docx")
    parser.add_argument("--include-jcr", action="store_true")
    args = parser.parse_args()
    content = json.loads(Path(args.content_json).read_text())
    errors = validate(content)
    if errors:
        raise SystemExit("\n".join(errors))
    skill_dir = Path(__file__).resolve().parents[1]
    build(content, args.out_docx, include_jcr=args.include_jcr, skill_dir=skill_dir)
    print(json.dumps({"out_docx": args.out_docx, "included": content.get("included_count"), "excluded": content.get("excluded_count")}, ensure_ascii=False))


if __name__ == "__main__":
    main()
